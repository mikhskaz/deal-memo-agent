"""Render a Memo to DOCX format using python-docx."""

from __future__ import annotations

import logging
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from models.deal import Memo, DealExtraction, SearchResult

logger = logging.getLogger(__name__)

SECTION_TITLES = {
    "executive_summary": "Executive Summary",
    "business_description": "Business Description",
    "market_opportunity": "Market Opportunity",
    "financial_overview": "Financial Overview",
    "key_risks": "Key Risks",
    "management_team": "Management Team",
    "diligence_questions": "Diligence Questions",
    "recommended_next_step": "Recommended Next Steps",
}

SECTION_ORDER = [
    "executive_summary",
    "business_description",
    "market_opportunity",
    "financial_overview",
    "key_risks",
    "management_team",
    "diligence_questions",
    "recommended_next_step",
]


def render_docx(
    memo: Memo,
    extraction: DealExtraction,
    sources: list[SearchResult],
    output_path: str | Path,
) -> Path:
    """Render a memo to a DOCX file.

    Creates a formatted Word document with:
    - AI draft disclaimer header
    - Company name and generation timestamp
    - All memo sections in standard order
    - Sources appendix
    """
    output_path = Path(output_path)
    doc = Document()

    # -- Document properties --
    doc.core_properties.author = "Deal Memo Agent (AI Draft)"
    doc.core_properties.subject = extraction.company_name or "Unknown Company"

    # -- Disclaimer header --
    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run(
        "AI-ASSISTED DRAFT — FOR INTERNAL REVIEW ONLY — NOT FOR DISTRIBUTION"
    )
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)  # Red

    # -- Title --
    company_name = extraction.company_name or "Deal"
    title = doc.add_heading(f"Investment Memo: {company_name}", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # -- Metadata line --
    meta = doc.add_paragraph()
    meta.add_run(f"Generated: {memo.generated_at.strftime('%Y-%m-%d %H:%M UTC')}").italic = True
    meta.add_run(f"  |  Model: {memo.model_used}").italic = True
    doc.add_paragraph()  # spacer

    # -- Sections --
    for section_id in SECTION_ORDER:
        section = memo.sections.get(section_id)
        if not section:
            continue

        doc.add_heading(SECTION_TITLES.get(section_id, section_id), level=1)

        # Parse markdown content into paragraphs
        for line in section.content.split("\n"):
            line = line.strip()
            if not line:
                continue

            # Handle markdown headings within section
            if line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("- ") or line.startswith("* "):
                p = doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("1. ") or line.startswith("2. ") or line.startswith("3. "):
                # Numbered list - strip the number prefix
                text = line.split(". ", 1)[1] if ". " in line else line
                p = doc.add_paragraph(text, style="List Number")
            else:
                # Handle bold markers
                if line.startswith("**") and line.endswith("**"):
                    p = doc.add_paragraph()
                    run = p.add_run(line.strip("*"))
                    run.bold = True
                else:
                    doc.add_paragraph(line)

    # -- Sources appendix --
    if sources:
        doc.add_page_break()
        doc.add_heading("Sources", level=1)
        for i, source in enumerate(sources, 1):
            p = doc.add_paragraph()
            p.add_run(f"[{i}] ").bold = True
            p.add_run(f"{source.title}\n")
            run = p.add_run(source.url)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x00, 0x00, 0xCC)
            p.add_run(f"\nQuery: {source.query}")

    doc.save(str(output_path))
    logger.info("DOCX saved to %s", output_path)
    return output_path
